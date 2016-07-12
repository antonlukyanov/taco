import os.path as path
import subprocess as sp
import re

from abc import ABC
from abc import abstractmethod

import jinja2 as jj

from docx import Document
from docx.shared import Emu
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import Mm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class BaseConverter(ABC):
    def __init__(self, ext, theme):
        self._ext = ext
        self._theme = theme

    @abstractmethod
    def convert(self, lists, out_filepath):
        pass


class TextConverter(BaseConverter):
    def __init__(self, ext, theme, templates_dir=None):
        super().__init__(ext, theme)
        self._path_prefix = path.join(self._ext, self._theme or '')
        self._templates_dir = templates_dir or path.join(path.dirname(__file__), 'templates')
        self._templates_path = path.join(self._templates_dir, self._path_prefix)
        self._env = jj.Environment(
            loader=jj.FileSystemLoader(self._templates_dir),
            trim_blocks=True,
            lstrip_blocks=True,
            comment_start_string='{--',
            comment_end_string='--}'
        )
        self._fallback_tpl_path = path.join(self._path_prefix, 'list.jinja2')
        self._layout = self._env.get_template('%s/layout.jinja2' % self._path_prefix)

    def _prepare_settings_hook(self, settings):
        pass

    def convert(self, tables, out_filepath):
        sections = []

        layout, layout_s = tables[0]['table'], tables[0]['settings']
        tables = tables[1:]

        allowed_width = layout_s.width - layout_s.margin_left\
                                       - layout_s.margin_right

        for data in tables:
            tp, table, settings = data['type'], data['table'], data['settings']
            self._prepare_settings_hook(settings)

            if settings['template'] == 'table':
                total_width = 0
                for col in settings['columns']:
                    width, unit = layout_s.parse_length(col['width'])
                    if unit == 'cm':
                        total_width += Cm(width).mm
                    else:
                        total_width += width

                if total_width > allowed_width:
                    raise RuntimeError('Total columns width of %.2f cm exceeded available width of'
                                       '%.2f cm.' % (Mm(total_width).cm, Mm(allowed_width).cm))

            # Loading template.
            tpl_name = settings['template'] + '.jinja2'
            # If template does not exist then fall back to list.jinja2.
            if not path.isfile(path.join(self._templates_path, tpl_name)):
                tpl_path = self._fallback_tpl_path
            else:
                tpl_path = path.join(self._path_prefix, tpl_name)
            tpl = self._env.get_template(tpl_path)
            # Rendering template.
            sections.append(tpl.render(tbl=table, settings=settings))

        with open(out_filepath, 'w') as f:
            f.write(self._layout.render(sections=sections, layout=layout, settings=layout_s))


class Latex2PdfConverter(TextConverter):
    def __init__(self, ext, theme, templates_dir=None):
        super().__init__(ext, theme, templates_dir)

    def convert(self, tables, out_filepath):
        filepath_tex = path.splitext(out_filepath)[0] + '.tex'
        super().convert(tables, filepath_tex)
        out_dir = path.dirname(out_filepath)
        if not out_dir:
            out_dir = path.curdir
        completed = sp.run(['pdflatex',
                            '-output-directory=%s' % out_dir,
                            filepath_tex])
        if completed.returncode == 0:
            print()
            print('Successfully converted to pdf.')
        else:
            raise RuntimeError('Could not convert to pdf.')

    def _prepare_settings_hook(self, settings):
        tex_cols = []
        if 'columns' in settings:
            for col in settings['columns']:
                if 'alignment' in col:
                    align = col['alignment']
                else:
                    align = 'l'

                if align not in ['l', 'c', 'r', 'p', 'm', 'b']:
                    raise RuntimeError('Unknown column alignment %s for column: %s.' % (align,
                                                                                       col['field']))
                if align in ['l', 'c', 'r'] and col['width']:
                    col_def = align.upper() + '{' + col['width'] + '}'
                elif align in ['p', 'm', 'b']:
                    col_def = align + '{' + col['width'] + '}'
                else:
                    col_def = align

                col['col_def'] = col_def
                tex_cols.append(col_def)

                if col['width']:
                    header_col_def = 'C{' + col['width'] + '}'
                else:
                    header_col_def = 'c'

                col['header_col_def'] = header_col_def

            settings['table_def'] = '|' + '|'.join(tex_cols) + '|'


class DocxConverter(BaseConverter):
    def __init__(self, ext, theme):
        super().__init__(ext, theme)

    def _setup_styles(self, doc):
        title = doc.styles.add_style('listconv.title', WD_STYLE_TYPE.PARAGRAPH)
        title_par = title.paragraph_format
        title_par.space_after = 0
        title_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.font.size = Pt(20)

        subtitle = doc.styles.add_style('listconv.subtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_par = subtitle.paragraph_format
        subtitle_par.space_after = Cm(0.6)
        subtitle_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle.font.size = Pt(14)

        h1 = doc.styles.add_style('listconv.h1', WD_STYLE_TYPE.PARAGRAPH)
        h1_par = h1.paragraph_format
        h1_par.space_after = Cm(0.6)
        h1.font.size = Pt(18)

        h2 = doc.styles.add_style('listconv.h2', WD_STYLE_TYPE.PARAGRAPH)
        h2_par = h2.paragraph_format
        h2_par.space_after = Cm(0.6)
        h2.font.size = Pt(16)

        indent = doc.styles.add_style('listconv.table_bottom_indent', WD_STYLE_TYPE.TABLE)
        indent_par = indent.paragraph_format
        indent_par.space_after = Cm(1)

        indent = doc.styles.add_style('listconv.bottom_indent', WD_STYLE_TYPE.PARAGRAPH)
        indent_par = indent.paragraph_format
        indent_par.space_after = Cm(1)

        indent = doc.styles.add_style('listconv.top_indent', WD_STYLE_TYPE.PARAGRAPH)
        indent_par = indent.paragraph_format
        indent_par.space_before = Cm(1)

        bc = doc.styles.add_style('listconv.bold_center', WD_STYLE_TYPE.PARAGRAPH)
        bc.font.bold = True
        bc.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        center = doc.styles.add_style('listconv.center', WD_STYLE_TYPE.PARAGRAPH)
        center.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _setup_page(self, doc, settings):
        section = doc.sections[0]
        section.page_width, section.page_height = (Mm(settings.width), Mm(settings.height))

        for side in ['left', 'right', 'top', 'bottom']:
            margin = Mm(getattr(settings, 'margin_%s' % side))
            setattr(section, '%s_margin' % side, margin)

        if settings.orientation == 'landscape':
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

    def convert(self, tables, out_filepath):
        doc = Document()

        layout_settings = tables[0]['settings']
        self._setup_styles(doc)
        self._setup_page(doc, layout_settings)

        sect = doc.sections[0]
        avail_width = Emu(sect.page_width - sect.left_margin - sect.right_margin).cm

        for data in tables:
            tp = data['type']
            table = data['table']

            if tp == 'layout':
                layout = table
                if layout['title']:
                    doc.add_paragraph(layout['title'], 'listconv.title')
                if layout['subtitle']:
                    doc.add_paragraph(layout['subtitle'], 'listconv.subtitle')
                if layout['description']:
                    doc.add_paragraph(layout['description'], 'listconv.bottom_indent')
                continue

            doc.add_paragraph(table.title, 'listconv.h2')
            if table.description:
                doc.add_paragraph(table.description)

            defs = data['settings']
            if defs['template'] == 'table' and defs['columns']:
                tbl = doc.add_table(
                    rows=0, cols=len(defs['columns']),
                    style='Table Grid')
                tbl.autofit = False

                total_width = Cm(0)
                for i, col in enumerate(defs['columns']):
                    m = re.match('^([0-9]+(\.[0-9]+)?)cm$', col['width'])
                    if m is not None:
                        col['width'] = Cm(float(m.group(1)))
                    else:
                        raise RuntimeError('Incorrect width "%s" for column "%s". Specify column '
                                           'width in form <X>cm, where <X> is number.'
                                           % (col['width'], col['name']))
                    total_width += col['width'].cm

                if total_width > avail_width:
                    raise RuntimeError('Total columns width of %.2f cm exceeded available width of'
                                       '%.2f cm for page.'
                                       % (total_width, avail_width))

                tbl.add_row()
                for i, col in enumerate(defs['columns']):
                    cell = tbl.cell(0, i)
                    cell.text = col['name'].strip()
                    cell.width = col['width']
                    cell.paragraphs[0].style = 'listconv.bold_center'
                    tbl.columns[i].width = col['width']

                for i, item in enumerate(table.rows):
                    row = tbl.add_row()
                    for j, col in enumerate(defs['columns']):
                        if col['field'] == 'INDEX':
                            row.cells[j].text = str(i+1)
                            row.cells[j].paragraphs[0].style = 'listconv.center'
                        else:
                            val = item[col['field']]
                            if val:
                                row.cells[j].text = str(val).strip()

                doc.add_paragraph('', 'listconv.bottom_indent')
            elif defs['template'] == 'list' and defs['items']:
                for it in table.rows:
                    vals = []
                    for field in defs['items']:
                        val = it[field].strip() or ''
                        if val:
                            vals.append(val)
                    doc.add_paragraph(', '.join(vals), style='List Number')

        doc.save(out_filepath)


# Tbd.
class HtmlConverter:
    pass
